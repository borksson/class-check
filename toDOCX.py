import json
from docx import Document
from datetime import datetime
from datetime import date
from datetime import timedelta  
from docx.shared import Pt

BUFFER = timedelta(days=1)

def JSONtoDOCX(JSONFile):
	document = Document();
	style = document.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = Pt(12)
	p = document.add_paragraph("CLASSCHECK:\nCLASSCHECK updated: " + date.today().strftime("%#m/%#d/%Y") + "\nDatafile updated: " + datetime.strptime(JSONFile[0]["updated"], '%Y-%m-%d %X.%f').strftime("%#m/%#d/%Y") + "\nBuffer: " + str(BUFFER))
	p.style = document.styles['Normal']
	for _class in JSONFile[0]["classes"]:
		s = ""
		for _assignment in _class["assignments"]:
			if(_assignment["dueDate"]!=''):
				dueDate = datetime.strptime(_assignment["dueDate"], '%Y-%m-%d %X')
			else:
				dueDate = datetime.today()
			if(dueDate.date() > date.today()):
				#TITLE (DUE DATE + BUFFER),
				#"2020-04-28 23:59:00
				s = s + _assignment["title"] + " (Due " + (dueDate - BUFFER).strftime("%#m/%#d/%Y, %H:%M:%S")  + ")"
				if(_assignment!=_class['assignments'][-1]):
					s = s + ", "
		p = document.add_paragraph(_class["className"] + ": " + s)
		p.style = document.styles['Normal']
	try:
		document.save("C://Users//craz0//Desktop//Files//Notes//CLASSCHECK.docx")
	except:
		print("Couldn't save document. Please close document.")

with open("dataFile.json") as f:
	dataFile = json.load(f)
f.close()
JSONtoDOCX(dataFile)
